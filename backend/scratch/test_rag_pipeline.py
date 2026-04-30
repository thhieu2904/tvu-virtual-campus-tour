"""
Quick integration test for the RAG pipeline.
Run with: conda activate tvu-tour && python scratch/test_rag_pipeline.py
"""
import httpx
import time
import sys

BASE = "http://localhost:8000/api"
ADMIN_KEY = "my-super-secret-admin-key"
HEADERS = {"X-Admin-Key": ADMIN_KEY}


def test_health():
    r = httpx.get(f"{BASE}/health")
    assert r.status_code == 200
    print(f"✅ Health: {r.json()['status']}")


def test_list_locations():
    r = httpx.get(f"{BASE}/locations")
    data = r.json()
    # API may return {locations: [...]} or [...]
    locs = data.get("locations", data) if isinstance(data, dict) else data
    print(f"✅ Locations: {len(locs)} found")
    for loc in locs[:3]:
        print(f"   - {loc['slug']:20s} {loc['name']}  (id: {loc['id'][:8]}...)")
    return locs


def test_ingest(location_id=None):
    """Upload a small test PDF to test the ingest pipeline."""
    # Create a minimal PDF with reportlab, or use a simple text-based approach
    # For now, let's create a simple DOCX since python-docx is available
    from docx import Document
    
    doc = Document()
    doc.add_heading("Giới thiệu Đại học Trà Vinh", level=1)
    doc.add_paragraph(
        "Đại học Trà Vinh (TVU) được thành lập năm 2006, tiền thân là Trường "
        "Cao đẳng Cộng đồng Trà Vinh. Trường tọa lạc tại số 126, Nguyễn Thiện Thành, "
        "Phường 5, Thành phố Trà Vinh, tỉnh Trà Vinh."
    )
    doc.add_heading("Ngành đào tạo", level=2)
    doc.add_paragraph(
        "Khoa Công nghệ Thông tin có 3 ngành đào tạo chính: "
        "Công nghệ Thông tin, Kỹ thuật Phần mềm, và Hệ thống Thông tin. "
        "Học phí trung bình khoảng 15-20 triệu đồng/năm."
    )
    doc.add_heading("Cơ sở vật chất", level=2)
    doc.add_paragraph(
        "TVU có hơn 20 phòng máy tính, thư viện hiện đại với hơn 50,000 đầu sách, "
        "và nhiều phòng thí nghiệm chuyên dụng."
    )
    
    import io
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    file_bytes = buf.read()
    
    print(f"\n📤 Uploading test DOCX ({len(file_bytes)} bytes)...")
    
    files = {"file": ("test_tvu_info.docx", file_bytes, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")}
    data = {"title": "Giới thiệu TVU (Test)"}
    if location_id:
        data["location_id"] = location_id
    
    r = httpx.post(f"{BASE}/admin/ingest", headers=HEADERS, files=files, data=data, timeout=30)
    print(f"   Response: {r.status_code} {r.json()}")
    
    if r.status_code != 202:
        print(f"❌ Ingest failed: {r.text}")
        return None
    
    doc_id = r.json()["document_id"]
    print(f"✅ Ingest started: document_id={doc_id}")
    
    # Poll status until ready
    for i in range(20):
        time.sleep(2)
        sr = httpx.get(f"{BASE}/admin/documents/{doc_id}/status", headers=HEADERS)
        status_data = sr.json()
        status = status_data["status"]
        print(f"   [{i+1}] Status: {status} (chunks: {status_data.get('chunk_count', 0)})")
        
        if status == "ready":
            print(f"✅ Document ready! {status_data['chunk_count']} chunks created")
            return doc_id
        elif status == "error":
            print(f"❌ Processing error: {status_data.get('error_message')}")
            return doc_id
    
    print("⏰ Timeout waiting for processing")
    return doc_id


def test_chat(location_id=None):
    """Test the chat endpoint with a RAG query."""
    print("\n💬 Testing chat...")
    
    body = {
        "message": "Khoa CNTT có mấy ngành đào tạo?",
        "location_id": location_id or "test",
    }
    
    r = httpx.post(f"{BASE}/chat", json=body, timeout=60)
    print(f"   Status: {r.status_code}")
    
    if r.status_code == 200:
        data = r.json()
        print(f"   Answer: {data.get('answer', 'N/A')[:200]}")
        print(f"   Sources: {len(data.get('sources', []))} chunks")
        print(f"   Time: {data.get('response_time_ms', 0)}ms")
        if data.get("sources"):
            for s in data["sources"][:2]:
                sim = s.get("similarity", 0)
                content = s.get("content", "")[:80]
                print(f"     - [{sim:.3f}] {content}...")
        print("✅ Chat OK!")
    else:
        print(f"❌ Chat failed: {r.text[:200]}")


def test_list_documents():
    """List all documents."""
    r = httpx.get(f"{BASE}/admin/documents", headers=HEADERS)
    data = r.json()
    print(f"\n📋 Documents: {data['total']} total")
    for doc in data.get("documents", []):
        print(f"   - [{doc['status']:10s}] {doc['title']} ({doc['chunk_count']} chunks)")


if __name__ == "__main__":
    print("=" * 60)
    print("RAG Pipeline Integration Test")
    print("=" * 60)
    
    test_health()
    locs = test_list_locations()
    
    # Use first location if available
    first_loc_id = locs[0]["id"] if locs else None
    
    # Test ingest
    doc_id = test_ingest(location_id=first_loc_id)
    
    # List documents
    test_list_documents()
    
    # Test chat
    if first_loc_id:
        test_chat(location_id=first_loc_id)
    
    print("\n" + "=" * 60)
    print("Done!")
